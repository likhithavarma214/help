"""
AccuKnox RFP Response Document Generator
Generates a professional .docx file with all required sections.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "AccuKnox_RFP_Response.docx")

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1B3A5C")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "EDF2F9")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_heading_numbered(doc, text, level):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(10)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


# ── Cover Page ────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AccuKnox")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("RFP Response Document")
    run2.font.size = Pt(24)
    run2.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.add_paragraph("")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Zero Trust Cloud-Native Application Protection Platform (CNAPP)")
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(4):
        doc.add_paragraph("")

    details = [
        "Prepared by: AccuKnox Inc.",
        "333 Ravenswood Ave, Menlo Park, CA 94025, USA",
        "support@accuknox.com | https://accuknox.com",
        "CONFIDENTIAL",
    ]
    for d in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(d)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


# ── Section 1 ─────────────────────────────────────────────────────────────────

def add_section_1(doc):
    add_heading_numbered(doc, "1. Company Introduction", 1)

    add_body(doc,
        "AccuKnox delivers a Zero Trust Security platform for AI, API, Application, Cloud, "
        "and Supply Chain Security. Incubated out of SRI International (Stanford Research Institute), "
        "AccuKnox holds seminal Zero Trust security patents and is backed by top-tier investors "
        "including National Grid Partners, Dolby Family Ventures, Dreamit Ventures, Avanta Ventures, "
        "and the 5G Open Innovation Lab."
    )

    add_body(doc,
        "AccuKnox has raised $15 million in seed funding and is headquartered in Menlo Park, California "
        "with engineering offices in Bengaluru and Chennai, India. The company is the creator of KubeArmor, "
        "an open-source runtime security engine with over 2 million downloads and 1,000+ GitHub stars, "
        "contributed to the CNCF ecosystem."
    )

    add_heading_numbered(doc, "Market Position", 2)
    bullets = [
        "Rated 4.7+ on Gartner Peer Insights for CNAPP",
        "Certified & accredited by NVIDIA, AWS Partner, NetApp, Nutanix",
        "Featured in Gartner, IBM Cloud, Oracle, and LF Edge",
        "Available on AWS, Azure, Oracle, Red Hat, Google Cloud, and Alibaba Cloud marketplaces",
        "10+ patents registered with the United States Patent & Trademark Office",
        "Supports 45+ global compliance frameworks",
        "Trusted by US DoD, Pure Storage, Sonesta, IDT Corporation, Prudent Insurance, and others",
    ]
    for b in bullets:
        add_bullet(doc, b)

    doc.add_page_break()


# ── Section 2 ─────────────────────────────────────────────────────────────────

def add_section_2(doc):
    add_heading_numbered(doc, "2. Platform Coverage", 1)

    add_body(doc,
        "The AccuKnox platform provides unified security from code to cloud to runtime, "
        "covering the following asset categories:"
    )

    headers = ["Asset Category", "Coverage Details"]
    rows = [
        ["Public Clouds", "AWS, Azure, GCP, Oracle Cloud, Alibaba Cloud"],
        ["Private Clouds", "OpenStack, OpenShift, VMware, Nutanix"],
        ["Kubernetes", "EKS, AKS, GKE, on-prem K8s, bare-metal, edge/IoT — 10+ K8s engines"],
        ["Virtual Machines", "EC2, Azure VMs, GCE, on-prem VMs, bare metal"],
        ["Containers & Serverless", "Docker, containerd, Podman, serverless functions"],
        ["APIs", "North-south and east-west API traffic, REST, gRPC, GraphQL"],
        ["AI/LLM Assets", "Hugging Face, OpenAI, TensorFlow, Ollama, managed & private models — 100+ models"],
        ["Application Pipelines", "CI/CD (GitHub Actions, GitLab, Jenkins, and 9+ platforms), IaC, SBOM"],
        ["Supply Chain", "Container registries (ECR, ACR, JFrog, Harbor, Quay, Docker Hub)"],
    ]
    add_styled_table(doc, headers, rows, col_widths=[5, 12])

    doc.add_page_break()


# ── Section 3 ─────────────────────────────────────────────────────────────────

def add_section_3(doc):
    add_heading_numbered(doc, "3. Deployment Model", 1)

    add_body(doc,
        "AccuKnox supports multiple deployment models to match operational, regulatory, and "
        "infrastructure needs. Core platform capabilities — scaling, high availability, and self-healing — "
        "remain consistent across all deployment modes."
    )

    headers = ["Model", "Description", "Key Characteristics"]
    rows = [
        ["SaaS (Fully Managed)", "AccuKnox manages the control plane on cloud infrastructure.",
         "Fastest to deploy; multi-tenant; AccuKnox handles upgrades, maintenance, and data retention. All features supported."],
        ["Managed OEM / MSSP", "AccuKnox operates the platform for partners delivering managed services.",
         "Partner stays customer-facing; AccuKnox maintains platform operations."],
        ["AWS Hybrid (Cloud + On-Prem)", "Control plane uses AWS managed services (S3, RDS) while workloads remain on-prem.",
         "Balances scalability with partial infrastructure control."],
        ["Full On-Prem / Air-Gapped", "Deployed entirely in customer environment with no shared infrastructure.",
         "Highest isolation; customer manages upgrades with AccuKnox support. AI CoPilot not included."],
    ]
    add_styled_table(doc, headers, rows, col_widths=[4, 5.5, 7.5])

    add_heading_numbered(doc, "Control Plane Architecture", 2)
    add_body(doc,
        "The AccuKnox control plane is fully deployed on Kubernetes, using a k8s-native model:"
    )
    arch_bullets = [
        "Stateless microservices with state externalized to shared databases (PostgreSQL, MongoDB, Vector DB, Graph DB)",
        "Horizontal scaling via Kubernetes HPA based on CPU, memory, and queue-depth metrics",
        "Kueue-based scheduling for fair tenant resource allocation and noisy-neighbor prevention",
        "SPIFFE-based identity for secure worker cluster onboarding",
        "Vault for dynamic secrets injection; k8s ConfigMaps for configuration management",
        "Multi-AZ deployment with pod anti-affinity and topology spread constraints",
        "Rolling updates with readiness/liveness probes for zero-downtime upgrades",
        "RTO: 6 hours | RPO: 24 hours",
    ]
    for b in arch_bullets:
        add_bullet(doc, b)

    add_body(doc, "Production uptime is publicly monitored at status.accuknox.com.")

    doc.add_page_break()


# ── Section 4 ─────────────────────────────────────────────────────────────────

def add_section_4(doc):
    add_heading_numbered(doc, "4. Modules Supported", 1)

    headers = ["Module", "Full Name", "Summary"]
    rows = [
        ["CSPM", "Cloud Security Posture Management",
         "Continuous misconfiguration detection, compliance mapping, and auto-remediation across multi-cloud environments."],
        ["KSPM", "Kubernetes Security Posture Management",
         "RBAC visualization, CIS benchmarking, cluster hardening, and identity entitlement management for Kubernetes."],
        ["ASPM", "Application Security Posture Management",
         "Unified SAST, SCA, DAST, IaC scanning, and SBOM analysis across the CI/CD pipeline."],
        ["CWPP", "Cloud Workload Protection Platform",
         "eBPF-based runtime protection, Zero Trust policy enforcement, micro-segmentation, and container forensics."],
        ["API Security", "API Discovery & Protection",
         "Runtime and static API discovery, OWASP API Top 10 coverage, traffic analysis, and access control."],
        ["AI Security", "AI-SPM & Red Teaming",
         "AI asset inventory, prompt firewall, automated red teaming, model sandboxing, and AI-GRC compliance."],
        ["Secrets Mgmt", "Secrets Management",
         "Centralized credential storage, dynamic secrets, encryption-as-a-service, HashiCorp Vault–compatible replacement."],
        ["SIEM", "Security Information & Event Management",
         "AI-driven threat detection, centralized log aggregation, real-time correlation, and SOC-optimized UI."],
        ["VM Security", "Virtual Machine Security",
         "Agentless and agent-based vulnerability scanning for VMs across cloud and on-prem environments."],
    ]
    add_styled_table(doc, headers, rows, col_widths=[3, 5.5, 8.5])

    doc.add_page_break()


# ── Section 5 ─────────────────────────────────────────────────────────────────

def add_section_5(doc):
    add_heading_numbered(doc, "5. Module Details", 1)

    # CSPM
    add_heading_numbered(doc, "5.1 CSPM — Cloud Security Posture Management", 2)
    add_body(doc,
        "AccuKnox CSPM continuously monitors cloud infrastructure across AWS, Azure, and GCP "
        "to detect misconfigurations, enforce compliance, and reduce cloud risk at scale."
    )
    add_heading_numbered(doc, "How It Works", 3)
    add_body(doc,
        "Agentless scanning connects to cloud provider APIs to inventory assets (compute, storage, "
        "network, IAM), evaluate configurations against security benchmarks, and generate prioritized "
        "findings with auto-remediation playbooks."
    )
    add_heading_numbered(doc, "Key Capabilities", 3)
    for b in [
        "Multi-cloud asset inventory and misconfiguration detection (EC2, S3, IAM, RDS, GKE, etc.)",
        "Compliance mapping to NIST, PCI-DSS, CIS Benchmarks, GDPR, SOC 2, HIPAA — 30+ frameworks",
        "Policy-as-Code with customizable YAML/JSON guardrails",
        "Cloud attack path analysis and prioritized network/identity exposure",
        "Alert prioritization with severity scoring and auto-fix playbooks",
        "Continuous drift detection and monitoring",
        "Integrated ticketing and notification channel integrations",
    ]:
        add_bullet(doc, b)
    add_heading_numbered(doc, "Differentiators", 3)
    for b in [
        "Combines static posture checks with runtime context (unlike scan-only tools)",
        "Deep Kubernetes visibility alongside cloud service scanning",
        "Zero Trust architecture support with least-privilege IAM behavioral insights",
        "Open-source integration (OpenSCAP, KubeArmor)",
        "Agentless + agent-based deployment options",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # KSPM
    add_heading_numbered(doc, "5.2 KSPM — Kubernetes Security Posture Management", 2)
    add_body(doc,
        "AccuKnox KSPM provides Kubernetes-specific security posture management with RBAC analytics, "
        "cluster hardening, and identity entitlement management (KIEM)."
    )
    add_heading_numbered(doc, "How It Works", 3)
    add_body(doc,
        "Connects to Kubernetes clusters (EKS, AKS, GKE, on-prem, edge) to scan configurations "
        "against CIS and NSA benchmarks, visualize RBAC relationships, and detect permission drift."
    )
    add_heading_numbered(doc, "Key Capabilities", 3)
    for b in [
        "Full-text search across all RBAC entities (service accounts, role bindings, etc.)",
        "Interactive graph visualization of user-permission-resource relationships",
        "Predefined query packs to highlight unnecessary privileges and orphaned accounts",
        "CIS benchmarking and NSA hardening guide compliance",
        "Change history tracking to identify risky modifications over time",
        "Namespace and workload security monitoring",
        "Compliance mapping to SOC 2, PCI-DSS, HIPAA",
    ]:
        add_bullet(doc, b)
    add_heading_numbered(doc, "Differentiators", 3)
    for b in [
        "First CNAPP tool to provide KSPM with KIEM out of the box",
        "eBPF-powered runtime visibility combined with posture checks",
        "Multi-cloud and edge-ready (EKS, AKS, GKE, bare-metal)",
        "Backed by KubeArmor and CNCF-native projects",
    ]:
        add_bullet(doc, b)

    # ASPM
    add_heading_numbered(doc, "5.3 ASPM — Application Security Posture Management", 2)
    add_body(doc,
        "AccuKnox ASPM integrates vulnerability management, SCA, SAST, DAST, and IaC scanning "
        "into a unified platform that prioritizes critical vulnerabilities from code to cloud."
    )
    add_heading_numbered(doc, "How It Works", 3)
    add_body(doc,
        "Embeds security checks into CI/CD pipelines (GitHub Actions, GitLab, Jenkins). "
        "Scans code repositories, container images, IaC templates, and runtime deployments. "
        "Correlates findings across code, pipeline, and runtime to eliminate false positives."
    )
    add_heading_numbered(doc, "Key Capabilities", 3)
    for b in [
        "SAST + SCA integration for proprietary and open-source code scanning",
        "IaC security scanning (Terraform, Helm, CloudFormation)",
        "CI/CD pipeline security with build-time gating",
        "DAST for web application and API endpoint testing",
        "SBOM analysis and supply chain risk assessment",
        "Runtime drift detection between deployment artifacts and live behavior",
        "Threat modeling and attack path visualization",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # CWPP
    add_heading_numbered(doc, "5.4 CWPP — Cloud Workload Protection Platform", 2)
    add_body(doc,
        "AccuKnox CWPP provides eBPF-based runtime security for containers, VMs, and bare-metal, "
        "using KubeArmor for inline Zero Trust policy enforcement."
    )
    add_heading_numbered(doc, "How It Works", 3)
    add_body(doc,
        "Deploys as a KubeArmor DaemonSet on worker nodes using eBPF and BPF LSM hooks. "
        "Auto-discovers application behavior (processes, file system access, network connections), "
        "generates Zero Trust policies, and enforces them inline — blocking attacks before damage occurs."
    )
    add_heading_numbered(doc, "Key Capabilities", 3)
    for b in [
        "Automated Zero Trust policy discovery and generation",
        "Inline mitigation against zero-day attacks (not just detect-and-respond)",
        "Application micro-segmentation at pod level",
        "Process, file, and network allowlisting",
        "Hardening policies based on MITRE, NIST, CIS, PCI-DSS, STIGs",
        "Container forensics and runtime telemetry",
        "Custom policy editor for fine-grained control (observe/audit/enforce modes)",
    ]:
        add_bullet(doc, b)
    add_heading_numbered(doc, "Differentiators", 3)
    for b in [
        "Inline prevention (not detect-and-respond) — blocks attacks at kernel level in real time",
        "Zero false positives with declarative allowlist approach",
        "Open-source KubeArmor foundation (CNCF sandbox project, 2M+ downloads)",
        "Supports Kubernetes, VMs, bare metal, and edge/IoT environments",
    ]:
        add_bullet(doc, b)

    # API Security
    add_heading_numbered(doc, "5.5 API Security", 2)
    add_body(doc,
        "AccuKnox API Security provides continuous API discovery, inventory, and protection "
        "across Kubernetes and microservices environments."
    )
    add_heading_numbered(doc, "Key Capabilities", 3)
    for b in [
        "Runtime API discovery using service mesh sidecars/proxies with OpenTelemetry export",
        "Static API analysis scanning code repos and API specs (OpenAPI, Swagger, WSDL)",
        "Shadow, zombie, and orphan API detection",
        "OWASP API Top 10 vulnerability identification",
        "North-south and east-west traffic visibility and segmentation",
        "PII/PHI sensitive data detection in headers and responses",
        "DoS attack detection and TLS security with eBPF XDP",
        "Access policy control and behavioral analytics",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # AI Security
    add_heading_numbered(doc, "5.6 AI Security", 2)
    add_body(doc,
        "AccuKnox AI Security addresses AI-specific threats across models, agents, datasets, "
        "and pipelines with posture management, red teaming, and runtime protection."
    )
    add_heading_numbered(doc, "AI Security Modules", 3)
    for b in [
        "AI-SPM (AI Security Posture Management) — asset inventory, risk scoring, lineage mapping",
        "AI Red Teaming — automated testing for injections, hallucinations, toxicity, bias",
        "Prompt Firewall — PII redaction, prompt injection defense, toxicity filtering, code execution prevention",
        "Model Sandboxing — network isolation, file system protection, process whitelisting, DNS filtering",
        "AI Detection & Response — real-time monitoring, anomaly detection, automated remediation",
        "AI-GRC — compliance with NIST AI, MITRE AI, AISCP, SOC, and other frameworks",
    ]:
        add_bullet(doc, b)
    add_heading_numbered(doc, "Supported Platforms", 3)
    add_body(doc,
        "Azure AI Foundry, Copilot Studio, AWS Bedrock, Hugging Face, OpenAI, TensorFlow, Ollama, "
        "and 100+ supported models across leading LLMs."
    )

    # Secrets Management
    add_heading_numbered(doc, "5.7 Secrets Management", 2)
    add_body(doc,
        "AccuKnox Secrets Management is a centralized, managed solution for credential security "
        "that serves as a drop-in replacement for HashiCorp Vault."
    )
    add_heading_numbered(doc, "Key Capabilities", 3)
    for b in [
        "Encrypted secrets storage with versioning",
        "Dynamic secrets for temporary access to AWS, Kubernetes, databases",
        "Data encryption as a service (PKI, Transit encryption via API)",
        "Identity-based auth (LDAP, OIDC, OKTA) with granular permissions",
        "Full audit logging of every interaction",
        "Multi-tenancy with per-tenant namespaces",
        "Direct integration with AccuKnox SIEM and CDR",
        "Deployable on-prem and in air-gapped environments",
    ]:
        add_bullet(doc, b)

    # SIEM
    add_heading_numbered(doc, "5.8 SIEM", 2)
    add_body(doc,
        "AccuKnox SIEM is a cloud-native, AI-first security information and event management "
        "platform designed for modern SOC teams."
    )
    add_heading_numbered(doc, "Key Capabilities", 3)
    for b in [
        "AI-driven threat detection with ML-based correlation (up to 80% noise reduction)",
        "Centralized log aggregation from Kubernetes, containers, cloud services, endpoints",
        "Flexible ingestion (Syslog, KubeArmor, CloudTrail, Azure Logs, threat intel feeds)",
        "Pre-built compliance reporting (SOC 2, PCI, HIPAA, GDPR)",
        "SOC-optimized dark-mode UI designed to reduce analyst fatigue",
        "SOAR integration for automated incident response playbooks",
        "10,000+ events/second processing, 100GB+ daily ingestion, sub-second search",
        "Hot/warm/cold data tiering for cost efficiency",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # VM Security
    add_heading_numbered(doc, "5.9 VM Security", 2)
    add_body(doc,
        "AccuKnox provides vulnerability scanning for virtual machines across cloud and on-prem "
        "environments using both agentless and agent-based approaches."
    )
    add_heading_numbered(doc, "Key Capabilities", 3)
    for b in [
        "Agentless scanning via Terraform scripts (creates VM snapshots for analysis)",
        "Agent-based scanning using knoxctl CLI for local file/folder analysis",
        "Continuous vulnerability assessment across multi-cloud VM fleets",
        "Integration with AccuKnox findings dashboard for unified risk visibility",
    ]:
        add_bullet(doc, b)


# ── Section 6 ─────────────────────────────────────────────────────────────────

def add_section_6(doc):
    add_heading_numbered(doc, "6. Integrations", 1)

    add_body(doc,
        "AccuKnox supports 30+ integrations enabling bi-directional sharing of security findings "
        "across the cloud security ecosystem."
    )

    add_heading_numbered(doc, "6.1 Notification Integrations", 2)
    headers = ["Platform", "Description"]
    rows = [
        ["Slack", "Real-time security alerts and finding notifications to Slack channels"],
        ["Microsoft Teams", "Alert delivery to Teams channels for SOC and DevOps collaboration"],
        ["PagerDuty", "Incident escalation and on-call routing for critical findings"],
        ["Email (SMTP)", "Configurable email notifications for findings and report delivery"],
        ["Webhook", "Generic webhook integration for custom notification workflows"],
    ]
    add_styled_table(doc, headers, rows, col_widths=[4, 13])

    add_heading_numbered(doc, "6.2 Ticketing Integrations", 2)
    rows2 = [
        ["Jira", "Automatic ticket creation with finding context, severity, and remediation guidance"],
        ["ServiceNow", "ITSM integration for incident management and change request workflows"],
    ]
    add_styled_table(doc, headers, rows2, col_widths=[4, 13])

    add_heading_numbered(doc, "6.3 SIEM & Log Management Integrations", 2)
    rows3 = [
        ["Splunk", "Log forwarding and security event integration via HEC"],
        ["Microsoft Sentinel", "Native Azure Sentinel integration for cloud-native SOC operations"],
        ["AccuKnox SIEM", "Built-in SIEM with native ingestion from all AccuKnox modules"],
        ["Syslog", "Standard syslog forwarding for compatibility with any SIEM platform"],
    ]
    add_styled_table(doc, headers, rows3, col_widths=[4, 13])

    add_heading_numbered(doc, "6.4 Code & CI/CD Integrations", 2)
    rows4 = [
        ["GitHub Actions", "Container scan, IaC scan, and SAST actions available on GitHub Marketplace"],
        ["GitLab CI", "Pipeline integration for SAST, SCA, and IaC scanning"],
        ["Jenkins", "Plugin-based integration for build-time security gating"],
        ["Checkmarx", "Import SAST findings from Checkmarx into AccuKnox"],
        ["SonarQube", "Code quality and security scan result aggregation"],
        ["Snyk", "SCA and container vulnerability findings integration"],
        ["Semgrep / Opengrep", "SAST and secrets scanning integration for source code repositories"],
        ["Fortify", "Static code analysis findings import"],
        ["Veracode", "Application security scan result aggregation"],
    ]
    add_styled_table(doc, headers, rows4, col_widths=[4, 13])

    doc.add_page_break()


# ── Section 7 ─────────────────────────────────────────────────────────────────

def add_section_7(doc):
    add_heading_numbered(doc, "7. Workflow Automation", 1)

    add_heading_numbered(doc, "7.1 Rules Engine", 2)
    add_body(doc,
        "The AccuKnox Rules Engine enables automated policy enforcement and response workflows "
        "across all security modules. Security teams can define condition-action rules that trigger "
        "automatically based on finding attributes."
    )
    add_heading_numbered(doc, "Capabilities", 3)
    for b in [
        "Configurable rules based on severity, resource type, compliance framework, or custom attributes",
        "Automated actions: create tickets, send notifications, trigger remediation playbooks, suppress findings",
        "Multi-condition rules with AND/OR logic for precise targeting",
        "Rule templates for common automation scenarios",
        "Audit trail of all rule executions and outcomes",
    ]:
        add_bullet(doc, b)

    add_heading_numbered(doc, "7.2 Ask AI Assistance (Ada)", 2)
    add_body(doc,
        "AccuKnox provides a Gen-AI powered security assistant (Ada) that helps users investigate "
        "findings, understand risks, and accelerate remediation."
    )
    add_heading_numbered(doc, "Capabilities", 3)
    for b in [
        "Natural language queries about security findings, policies, and compliance status",
        "AI-generated remediation guidance with step-by-step instructions",
        "Context-aware analysis that correlates findings across modules",
        "Interactive exploration of cloud attack paths and risk relationships",
        "Available within the AccuKnox portal for SaaS deployments",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()


# ── Section 8 ─────────────────────────────────────────────────────────────────

def add_section_8(doc):
    add_heading_numbered(doc, "8. SLA, Support Coverage, and Resources", 1)

    add_heading_numbered(doc, "8.1 Support Tiers", 2)
    headers = ["Support Level", "Support Hours"]
    rows = [
        ["Silver (Standard)", "12x5 (6:00 AM – 6:00 PM PT, Mon-Fri, excluding US holidays)"],
        ["Gold (Premium)", "12x7 (Extended Phone Support)"],
        ["Platinum (PremiumPlus)", "24x7 (Full Round-the-Clock Support)"],
    ]
    add_styled_table(doc, headers, rows, col_widths=[5, 12])

    add_heading_numbered(doc, "8.2 Additional Support Features", 2)
    headers2 = ["Feature", "Silver (Standard)", "Gold (Premium)", "Platinum (PremiumPlus)"]
    rows2 = [
        ["Training", "Standard Training", "Standard Training + 2 Student Admin Webinar Series",
         "Live Training + 5 Student Premium Training Enrollments"],
        ["Phone Support", "12x5", "12x7", "24x7"],
        ["Online Support", "✓", "✓", "✓"],
        ["Adoption Services", "None", "None", "Named Customer Success Manager"],
    ]
    add_styled_table(doc, headers2, rows2, col_widths=[4, 4, 4.5, 4.5])

    add_heading_numbered(doc, "8.3 SLA Response Times", 2)
    headers3 = ["Severity Level", "Silver (Standard)", "Gold (Premium)", "Platinum (PremiumPlus)"]
    rows3 = [
        ["P1 (Critical)", "First Response: 4 Hours\nUpdates: 1 Business Day",
         "First Response: 1 Hour\nUpdates: 2 Hours", "First Response: 1 Hour\nUpdates: 2 Hours"],
        ["P2 (High)", "First Response: 1 Business Day\nUpdates: 2 Business Days",
         "First Response: 2 Hours\nUpdates: 8 Hours", "First Response: 2 Hours\nUpdates: 8 Hours"],
        ["P3 (Medium)", "First Response: 2 Business Days\nUpdates: 3 Business Days",
         "First Response: 2 Hours\nUpdates: 48 Hours", "First Response: 2 Hours\nUpdates: 48 Hours"],
        ["P4 (Low)", "First Response: 2 Business Days\nUpdates: 3 Business Days",
         "First Response: 8 Hours\nUpdates: 48 Hours", "First Response: 8 Hours\nUpdates: 48 Hours"],
    ]
    add_styled_table(doc, headers3, rows3, col_widths=[3.5, 4.5, 4.5, 4.5])

    doc.add_page_break()

    add_heading_numbered(doc, "8.4 Severity Levels & Definitions", 2)
    headers4 = ["Severity Level", "Definition", "Examples"]
    rows4 = [
        ["P1 (Critical)",
         "AccuKnox is down and inaccessible. Severe service failure or degradation affecting multiple users.",
         "Users cannot access a business-critical application. Consistent 'page not found' errors prevent login."],
        ["P2 (High)",
         "Partial service failure or mild degradation. Some, but not all, business resources are accessible.",
         "Admin console write-access issue. Users experience slow access, occasional errors. Bug causing significant impact."],
        ["P3 (Medium)",
         "Minor service impact, affecting individual users or non-critical third-party applications.",
         "One user is unable to access an application. Difficulty integrating new business applications."],
        ["P4 (Low)",
         "Minor impact or feature enhancement request.",
         "How-to inquiries. Feature enhancement requests."],
    ]
    add_styled_table(doc, headers4, rows4, col_widths=[3, 6.5, 7.5])

    add_heading_numbered(doc, "8.5 Escalation Matrix", 2)
    headers5 = ["Severity Level", "Personnel To Contact", "Email"]
    rows5 = [
        ["L1 (Low)", "Udit, Aditya — Sr Solution Engineers @AccuKnox",
         "udit@accuknox.com, adityaraj@accuknox.com"],
        ["L2 (Medium)", "Gaurav — Product Manager @AccuKnox", "gaurav.mishra@accuknox.com"],
        ["L3 (Critical)", "Rahul Jadhav — CTO @AccuKnox", "r@accuknox.com"],
    ]
    add_styled_table(doc, headers5, rows5, col_widths=[3, 6.5, 7.5])

    add_heading_numbered(doc, "8.6 Contact Methods", 2)
    add_body(doc, "You can reach our support team through the following methods, listed in order of preference:")
    add_bullet(doc, "Send an email directly to support@accuknox.com", bold_prefix="Email: ")
    add_bullet(doc,
        "Raise a support ticket via our customer portal: https://accu-knox.atlassian.net/servicedesk/customer/portal/1",
        bold_prefix="Jira Service Desk: ")
    add_body(doc, "Our support team will respond to your ticket within 24 business hours.")

    add_heading_numbered(doc, "8.7 Support Resources", 2)
    headers6 = ["Resource", "URL"]
    rows6 = [
        ["Support Email", "support@accuknox.com"],
        ["Help Portal", "https://help.accuknox.com"],
        ["Status Page", "https://status.accuknox.com"],
        ["Jira Service Desk", "https://accu-knox.atlassian.net/servicedesk/customer/portal/1"],
    ]
    add_styled_table(doc, headers6, rows6, col_widths=[5, 12])

    add_heading_numbered(doc, "8.8 Recovery Objectives", 2)
    headers7 = ["Metric", "Commitment", "Scope"]
    rows7 = [
        ["RTO (Recovery Time Objective)", "6 hours",
         "Restoration of core control plane services (policy management, telemetry, dashboards, enforcement)."],
        ["RPO (Recovery Point Objective)", "24 hours",
         "Control plane config, security policies, tenant config, audit logs, platform state."],
    ]
    add_styled_table(doc, headers7, rows7, col_widths=[4.5, 3, 9.5])

    add_body(doc,
        "Note: Runtime enforcement on customer workloads continues independently via deployed agents, "
        "ensuring fail-secure behavior even during control plane recovery."
    )


# ── Footer ────────────────────────────────────────────────────────────────────

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("AccuKnox Inc. | CONFIDENTIAL | support@accuknox.com | accuknox.com")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_cover_page(doc)
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)
    add_section_4(doc)
    add_section_5(doc)
    add_section_6(doc)
    add_section_7(doc)
    add_section_8(doc)
    add_footer(doc)

    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
