"""
AccuKnox RFP Response Document Generator — v2
Generates a professional .docx file with all required sections.
Includes: support matrices, architecture details, CDR, malware scanning,
container registry scanning, use cases, and help doc resource links.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "AccuKnox_RFP_Response.docx")
HELP_BASE = "https://help.accuknox.com"
WEB_BASE = "https://accuknox.com"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1B3A5C")
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "EDF2F9")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def h1(doc, text):
    doc.add_heading(text, level=1)

def h2(doc, text):
    doc.add_heading(text, level=2)

def h3(doc, text):
    doc.add_heading(text, level=3)

def body(doc, text):
    p = doc.add_paragraph(text)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def link_line(doc, label, url):
    """Add a line with a bold label and a URL."""
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(9)
    run2 = p.add_run(url)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def resource_block(doc, links):
    """Add a block of resource links. links = [(label, url), ...]"""
    p = doc.add_paragraph()
    run = p.add_run("Resources & Documentation:")
    run.bold = True
    run.font.size = Pt(9)
    for label, url in links:
        p2 = doc.add_paragraph(style="List Bullet")
        run2 = p2.add_run(f"{label}: {url}")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)


# ── Cover Page ────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AccuKnox")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("RFP Response Document")
    run2.font.size = Pt(24)
    run2.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.add_paragraph("")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Zero Trust Cloud-Native Application Protection Platform (CNAPP)")
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(4):
        doc.add_paragraph("")
    for d in [
        "Prepared by: AccuKnox Inc.",
        "333 Ravenswood Ave, Menlo Park, CA 94025, USA",
        "support@accuknox.com | https://accuknox.com",
        "CONFIDENTIAL",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(d)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()


# ── Section 1: Company Introduction ──────────────────────────────────────────

def add_section_1(doc):
    h1(doc, "1. Company Introduction")
    body(doc,
        "AccuKnox delivers a Zero Trust Security platform for AI, API, Application, Cloud, "
        "and Supply Chain Security. Incubated out of SRI International (Stanford Research Institute), "
        "AccuKnox holds seminal Zero Trust security patents and is backed by top-tier investors "
        "including National Grid Partners, Dolby Family Ventures, Dreamit Ventures, Avanta Ventures, "
        "and the 5G Open Innovation Lab."
    )
    body(doc,
        "AccuKnox has raised $15 million in seed funding and is headquartered in Menlo Park, California "
        "with engineering offices in Bengaluru and Chennai, India. The company is the creator of KubeArmor, "
        "an open-source runtime security engine with over 2 million downloads and 1,000+ GitHub stars, "
        "contributed to the CNCF ecosystem."
    )
    h2(doc, "Market Position")
    for b in [
        "Rated 4.7+ on Gartner Peer Insights for CNAPP",
        "Certified & accredited by NVIDIA, AWS Partner, NetApp, Nutanix",
        "Featured in Gartner, IBM Cloud, Oracle, and LF Edge",
        "Available on AWS, Azure, Oracle, Red Hat, Google Cloud, and Alibaba Cloud marketplaces",
        "10+ patents registered with the United States Patent & Trademark Office",
        "Supports 45+ global compliance frameworks",
        "Trusted by US DoD, Pure Storage, Sonesta, IDT Corporation, Prudent Insurance, and others",
    ]:
        bullet(doc, b)
    resource_block(doc, [
        ("Enterprise Architecture", f"{HELP_BASE}/getting-started/accuknox-arch/"),
        ("Control Plane Architecture", f"{HELP_BASE}/resources/control-plane-architecture/"),
        ("Deployment Models", f"{HELP_BASE}/getting-started/deployment-models/"),
    ])
    doc.add_page_break()


# ── Section 2: Platform Coverage ─────────────────────────────────────────────

def add_section_2(doc):
    h1(doc, "2. Platform Coverage")
    body(doc,
        "The AccuKnox platform provides unified security from code to cloud to runtime, "
        "covering the following asset categories:"
    )
    add_styled_table(doc,
        ["Asset Category", "Coverage Details"],
        [
            ["Public Clouds", "AWS, Azure, GCP, Oracle Cloud, Alibaba Cloud"],
            ["Private Clouds", "OpenStack, OpenShift, VMware, Nutanix"],
            ["Kubernetes", "EKS, AKS, GKE, on-prem K8s, bare-metal, edge/IoT — 10+ K8s engines"],
            ["Virtual Machines", "EC2, Azure VMs, GCE, on-prem VMs, bare metal (Linux + Windows)"],
            ["Containers & Serverless", "Docker, containerd, Podman, serverless functions"],
            ["Container Registries", "ECR, ACR, GAR, Harbor, DockerHub, JFrog, Quay, Nexus — 10 registries"],
            ["APIs", "North-south and east-west API traffic, REST, gRPC, GraphQL"],
            ["AI/LLM Assets", "Hugging Face, OpenAI, TensorFlow, Ollama, Azure AI, AWS Bedrock — 100+ models"],
            ["Application Pipelines", "CI/CD (GitHub Actions, GitLab, Jenkins, Azure DevOps, Bitbucket, CircleCI, Harness, GCP Cloud Build, AWS CodePipeline)"],
            ["Supply Chain", "SBOM analysis, IaC scanning, secrets scanning, SARIF findings import"],
        ],
        col_widths=[5, 12]
    )

    h2(doc, "Cloud Asset Coverage (CSPM)")
    body(doc, "AccuKnox CSPM scans 150+ distinct cloud asset types across AWS, Azure, and GCP, organized into 22+ categories:")
    for b in [
        "Compute: EC2, Azure VMs, GCP Compute Engine instances",
        "Storage: S3, Azure Storage Accounts, GCP Storage Buckets",
        "Networking: VPCs, Security Groups, Load Balancers, CDN",
        "IAM: Users, Roles, Policies, Service Accounts, MFA",
        "Database: RDS, Azure SQL, GCP Cloud SQL, DynamoDB, CosmosDB",
        "Serverless: Lambda, Azure Functions, Cloud Functions",
        "AI/ML: SageMaker, Azure ML, GCP Vertex AI, Bedrock",
        "Containers: ECS, AKS, GKE, EKS, Container Registries",
    ]:
        bullet(doc, b)
    resource_block(doc, [
        ("Full CSPM Asset List", f"{HELP_BASE}/support-matrix/assets-list/"),
        ("Cloud Regions Support", f"{HELP_BASE}/support-matrix/cloud-regions/"),
        ("Public Cloud Support", f"{HELP_BASE}/support-matrix/public-cloud/"),
        ("Private Cloud Support", f"{HELP_BASE}/support-matrix/private-cloud/"),
    ])
    doc.add_page_break()


# ── Section 3: Deployment Model ──────────────────────────────────────────────

def add_section_3(doc):
    h1(doc, "3. Deployment Model")
    body(doc,
        "AccuKnox supports multiple deployment models to match operational, regulatory, and "
        "infrastructure needs. Core platform capabilities — scaling, high availability, and self-healing — "
        "remain consistent across all deployment modes."
    )
    add_styled_table(doc,
        ["Model", "Description", "Key Characteristics"],
        [
            ["SaaS (Fully Managed)", "AccuKnox manages the control plane on cloud infrastructure.",
             "Fastest to deploy; multi-tenant; AccuKnox handles upgrades, maintenance, and data retention. All features supported."],
            ["Managed OEM / MSSP", "AccuKnox operates the platform for partners delivering managed services.",
             "Partner stays customer-facing; AccuKnox maintains platform operations."],
            ["AWS Hybrid (Cloud + On-Prem)", "Control plane uses AWS managed services (S3, RDS) while workloads remain on-prem.",
             "Balances scalability with partial infrastructure control."],
            ["Full On-Prem / Air-Gapped", "Deployed entirely in customer environment with no shared infrastructure.",
             "Highest isolation; customer manages upgrades with AccuKnox support. AI CoPilot not included."],
        ],
        col_widths=[4, 5.5, 7.5]
    )

    h2(doc, "Control Plane Architecture")
    body(doc, "The AccuKnox control plane is fully deployed on Kubernetes, using a k8s-native model:")
    for b in [
        "Stateless microservices with state externalized to shared databases (PostgreSQL, MongoDB, Vector DB, Graph DB)",
        "Horizontal scaling via Kubernetes HPA based on CPU, memory, and queue-depth metrics",
        "Kueue-based scheduling for fair tenant resource allocation and noisy-neighbor prevention",
        "SPIFFE-based identity for secure worker cluster onboarding",
        "Vault for dynamic secrets injection; k8s ConfigMaps for configuration management",
        "Multi-AZ deployment with pod anti-affinity and topology spread constraints",
        "Rolling updates with readiness/liveness probes for zero-downtime upgrades",
        "RTO: 6 hours | RPO: 24 hours",
    ]:
        bullet(doc, b)

    h2(doc, "Interaction Model: Customer Plane ↔ AccuKnox Control Plane")
    add_styled_table(doc,
        ["Feature", "Deployment Model", "Onboarding", "Permissions", "Playbook Hosted"],
        [
            ["CSPM (Org Units)", "Agentless", "Cross-tenant roles", "Read + Security Audit", "AccuKnox Control Plane"],
            ["CSPM (Individual Accounts)", "Agentless", "Shared key/id", "Read + Security Audit", "AccuKnox Control Plane"],
            ["CDR", "Agentless", "CloudFormation / Terraform", "Read logs + Create Lambda + S3", "Customer Cloud Env"],
            ["KSPM Scans", "Agentless", "Helm job deployment", "Read across namespaces", "Kubernetes Jobs"],
            ["Runtime Security (CWPP)", "Agent-based (KubeArmor)", "Helm DaemonSet", "eBPF on worker nodes", "K8s DaemonSet"],
            ["VM Vulnerability Scan", "Agentless", "Terraform script", "Create VM snapshot + scan", "Customer hosted VM"],
            ["VM Vulnerability Scan", "Agent-based", "knoxctl CLI", "File/folder scan", "knoxctl CLI tool"],
            ["SAST / IaC / Secrets", "Agentless", "Collectors-based", "Read code repos", "AccuKnox or customer hosted"],
            ["DAST", "Collectors model", "Credential config", "Read/Probe endpoints", "AccuKnox or customer hosted"],
            ["AI Red Teaming (Cloud)", "Agentless", "Same as cloud accounts", "Read + security audit", "AccuKnox hosted"],
            ["Prompt Firewall", "SDK / AI Gateway", "SDK integration or LiteLLM/BiFrost", "Prompt/response data", "Customer code or AccuKnox"],
        ],
        col_widths=[3.5, 2.5, 3.5, 3.5, 4]
    )
    resource_block(doc, [
        ("On-Prem Installation Guide", f"{HELP_BASE}/getting-started/on-prem-installation-guide/"),
        ("Managed Installation (EKS/AKS/GKE)", f"{HELP_BASE}/getting-started/onprem-eks-aks-gke/"),
        ("Single Node Installation", f"{HELP_BASE}/getting-started/on-prem-single-node-installation/"),
    ])
    doc.add_page_break()


# ── Section 4: Modules Supported ─────────────────────────────────────────────

def add_section_4(doc):
    h1(doc, "4. Modules Supported")
    add_styled_table(doc,
        ["Module", "Full Name", "Summary"],
        [
            ["CSPM", "Cloud Security Posture Management",
             "Continuous misconfiguration detection, compliance mapping, and auto-remediation across multi-cloud environments."],
            ["KSPM", "Kubernetes Security Posture Management",
             "RBAC visualization, CIS benchmarking, cluster hardening, and identity entitlement management for Kubernetes."],
            ["ASPM", "Application Security Posture Management",
             "Unified SAST, SCA, DAST, IaC scanning, and SBOM analysis across the CI/CD pipeline."],
            ["CWPP", "Cloud Workload Protection Platform",
             "eBPF-based runtime protection, Zero Trust policy enforcement, micro-segmentation, and container forensics."],
            ["CDR", "Cloud Detection & Response",
             "Real-time cloud log analysis (CloudTrail, Azure Audit, GCP Audit), automated remediation via CI/CD workflows."],
            ["API Security", "API Discovery & Protection",
             "Runtime and static API discovery, OWASP API Top 10 coverage, traffic analysis, and access control."],
            ["AI Security", "AI-SPM & Red Teaming",
             "AI asset inventory, prompt firewall, automated red teaming, model sandboxing, and AI-GRC compliance."],
            ["Secrets Mgmt", "Secrets Management",
             "Centralized credential storage, dynamic secrets, encryption-as-a-service, HashiCorp Vault–compatible replacement."],
            ["SIEM", "Security Information & Event Management",
             "AI-driven threat detection, centralized log aggregation, real-time correlation, and SOC-optimized UI."],
            ["VM Security", "Virtual Machine & Host Security",
             "Agentless + agent-based vulnerability scanning, malware detection (ClamAV), compliance benchmarking for Linux and Windows VMs."],
            ["Registry Scan", "Container Registry Scanning",
             "Vulnerability, secrets, and layer analysis for images across 10 container registries (ECR, ACR, GAR, Harbor, DockerHub, etc.)."],
        ],
        col_widths=[2.5, 5, 9.5]
    )
    doc.add_page_break()


# ── Section 5: Module Details ────────────────────────────────────────────────

def add_section_5(doc):
    h1(doc, "5. Module Details")

    # ── 5.1 CSPM ──
    h2(doc, "5.1 CSPM — Cloud Security Posture Management")
    body(doc,
        "AccuKnox CSPM continuously monitors cloud infrastructure across AWS, Azure, GCP, and Oracle "
        "to detect misconfigurations, enforce compliance, and reduce cloud risk at scale."
    )
    h3(doc, "How It Works")
    body(doc,
        "Agentless scanning connects to cloud provider APIs to inventory assets (compute, storage, "
        "network, IAM), evaluate configurations against 30+ compliance frameworks, and generate prioritized "
        "findings with auto-remediation playbooks. Drift detection monitors for configuration changes that "
        "violate established baselines."
    )
    h3(doc, "Support Matrix")
    add_styled_table(doc,
        ["Cloud Provider", "Asset Categories", "Compliance Frameworks"],
        [
            ["AWS", "EC2, S3, IAM, RDS, VPC, Lambda, CloudTrail, SageMaker, etc.", "CIS, NIST 800-53, PCI-DSS, SOC 2, HIPAA, GDPR"],
            ["Azure", "VMs, Storage Accounts, SQL, NSGs, AKS, Key Vault, etc.", "CIS, NIST, PCI-DSS, SOC 2, HIPAA, GDPR"],
            ["GCP", "Compute, Storage, IAM, Cloud SQL, GKE, BigQuery, etc.", "CIS, NIST, PCI-DSS, SOC 2, HIPAA"],
            ["Oracle", "Compute, Object Storage, Identity, VCN, etc.", "CIS, NIST"],
        ],
        col_widths=[3, 8, 6]
    )
    h3(doc, "Key Capabilities")
    for b in [
        "150+ cloud asset types scanned across AWS, Azure, GCP",
        "Multi-cloud asset inventory with automated discovery",
        "Compliance mapping to 30+ frameworks (NIST, CIS, PCI-DSS, SOC 2, HIPAA, GDPR, STIGs)",
        "Cloud attack path analysis and prioritized network/identity exposure",
        "Policy-as-Code with customizable YAML/JSON guardrails",
        "Continuous drift detection and baseline monitoring",
        "Alert prioritization with severity scoring and auto-fix playbooks",
        "Integrated ticketing (Jira, ServiceNow) and notification channels (Slack, Email)",
    ]:
        bullet(doc, b)
    h3(doc, "Use Cases")
    bullet(doc, "Detecting and remediating publicly exposed S3 buckets, open security groups, and over-permissive IAM roles", bold_prefix="Cloud Misconfiguration Remediation: ")
    bullet(doc, "Continuous GRC compliance monitoring with audit-ready report generation for SOC 2, PCI-DSS, HIPAA", bold_prefix="Compliance Audit Readiness: ")
    resource_block(doc, [
        ("CSPM Overview", f"{HELP_BASE}/use-cases/cspm/"),
        ("AWS Onboarding", f"{HELP_BASE}/how-to/aws-onboarding/"),
        ("Azure Onboarding", f"{HELP_BASE}/how-to/azure-onboarding/"),
        ("GCP Onboarding", f"{HELP_BASE}/how-to/gcp-onboarding/"),
        ("Asset Inventory", f"{HELP_BASE}/use-cases/asset-inventory/"),
        ("CSPM Asset Support List", f"{HELP_BASE}/support-matrix/assets-list/"),
    ])
    doc.add_page_break()

    # ── 5.2 KSPM ──
    h2(doc, "5.2 KSPM — Kubernetes Security Posture Management")
    body(doc,
        "AccuKnox KSPM provides Kubernetes-specific security posture management with RBAC analytics, "
        "cluster hardening, CIS benchmarking, and identity entitlement management (KIEM)."
    )
    h3(doc, "How It Works")
    body(doc,
        "Connects to Kubernetes clusters via agentless helm jobs to scan configurations "
        "against CIS and NSA benchmarks, visualize RBAC relationships via interactive graphs, "
        "detect permission drift, and enforce admission control policies via KnoxGuard."
    )
    h3(doc, "Support Matrix")
    add_styled_table(doc,
        ["K8s Engine", "Deployment Type", "Features Supported"],
        [
            ["Amazon EKS", "Managed", "CIS Benchmarks, KSPM, KIEM, Admission Controller"],
            ["Azure AKS", "Managed", "CIS Benchmarks, KSPM, KIEM, Admission Controller"],
            ["Google GKE", "Managed", "CIS Benchmarks, KSPM, KIEM, Admission Controller"],
            ["On-Prem Kubernetes", "Self-managed", "CIS Benchmarks, KSPM, KIEM, Admission Controller"],
            ["OpenShift", "Self-managed", "CIS Benchmarks, KSPM, KIEM"],
            ["Rancher / k3s", "Self-managed", "CIS Benchmarks, KSPM, KIEM"],
            ["Bare Metal", "Self-managed", "CIS Benchmarks, KSPM, KIEM"],
            ["Edge / IoT", "Self-managed", "CIS Benchmarks, KSPM"],
        ],
        col_widths=[4, 3.5, 9.5]
    )
    h3(doc, "Key Capabilities")
    for b in [
        "KIEM: Full-text search across RBAC entities, interactive graph visualization, critical query packs",
        "CIS Kubernetes Benchmark scanning with guided remediation",
        "Admission Controller (KnoxGuard) for supply chain attack prevention",
        "Pod Security Admission Control enforcement",
        "Cluster misconfiguration scanning and namespace security monitoring",
        "Change history tracking for risky RBAC modifications",
    ]:
        bullet(doc, b)
    h3(doc, "Use Cases")
    bullet(doc, "Visualizing over-permissive RBAC bindings and enforcing least-privilege across namespaces", bold_prefix="RBAC Governance: ")
    bullet(doc, "Blocking untrusted container images at deploy-time via KnoxGuard admission controller", bold_prefix="Supply Chain Protection: ")
    resource_block(doc, [
        ("KSPM Overview", f"{HELP_BASE}/use-cases/kspm/"),
        ("KIEM", f"{HELP_BASE}/use-cases/kiem/"),
        ("CIS Benchmarking", f"{HELP_BASE}/how-to/cis-benchmarking/"),
        ("Admission Controller", f"{HELP_BASE}/use-cases/admission-controller-knoxguard/"),
        ("KubeArmor Support Matrix", f"{HELP_BASE}/support-matrix/kubearmor-support-matrix/"),
    ])

    # ── 5.3 ASPM ──
    h2(doc, "5.3 ASPM — Application Security Posture Management")
    body(doc,
        "AccuKnox ASPM integrates vulnerability management, SCA, SAST, DAST, IaC scanning, "
        "secrets scanning, and SBOM analysis into a unified platform that prioritizes critical "
        "vulnerabilities from code to cloud."
    )
    h3(doc, "How It Works")
    body(doc,
        "Embeds security checks into CI/CD pipelines via workflow plugins or native integrations. "
        "Scans code repositories, container images, IaC templates, and runtime deployments. "
        "Correlates findings across code, pipeline, and runtime to eliminate false positives. "
        "EPSS scoring prioritizes exploitable vulnerabilities."
    )
    h3(doc, "CI/CD Support Matrix")
    add_styled_table(doc,
        ["Platform", "SAST", "DAST", "IaC Scan", "Container Scan", "Secret Scan"],
        [
            ["GitHub Actions", "✓", "✓", "✓", "✓", "✓"],
            ["GitLab CI/CD", "✓", "✓", "✓", "✓", "✓"],
            ["Jenkins", "✓", "✓", "✓", "✓", "✓"],
            ["Azure DevOps", "✓", "✓", "✓", "✓", "✓"],
            ["Bitbucket", "✓", "✓", "✓", "✓", "✓"],
            ["CircleCI", "✓", "✓", "✓", "✓", "✓"],
            ["AWS CodePipeline", "✓", "✓", "✓", "✓", "✓"],
            ["GCP Cloud Build", "✓", "✓", "✓", "✓", "—"],
            ["Harness", "✓", "✓", "✓", "✓", "—"],
            ["Bamboo CI", "✓", "✓", "✓", "✓", "✓"],
        ],
        col_widths=[3.5, 2, 2, 2, 2.5, 2.5]
    )
    h3(doc, "Key Capabilities")
    for b in [
        "SAST via Semgrep, Opengrep, SonarQube, Checkmarx, Fortify, Veracode integration",
        "SCA for open-source dependency risk and license compliance",
        "DAST with authenticated and MFA-enabled scan support",
        "IaC scanning for Terraform, Helm, CloudFormation, ARM templates",
        "Container image scanning in CI/CD pipelines and registries",
        "SBOM generation and supply chain risk assessment",
        "EPSS scoring for exploit probability-based vulnerability prioritization",
        "SARIF findings import for tool-agnostic result aggregation",
        "Scheduled and on-demand ASPM reports",
    ]:
        bullet(doc, b)
    h3(doc, "Use Cases")
    bullet(doc, "Embedding SAST/SCA/container scanning in GitHub Actions to block vulnerable builds before merge", bold_prefix="Shift-Left Pipeline Security: ")
    bullet(doc, "Scanning Terraform/Helm templates for misconfigurations before cloud deployment", bold_prefix="IaC Security Gates: ")
    resource_block(doc, [
        ("ASPM Overview", f"{HELP_BASE}/use-cases/aspm/"),
        ("CI/CD Support Matrix", f"{HELP_BASE}/support-matrix/cicd-support-matrix/"),
        ("DevSecOps Guide", f"{HELP_BASE}/getting-started/devsecops/"),
        ("Container Scan Use Case", f"{HELP_BASE}/use-cases/container-scan/"),
        ("IaC Scan", f"{HELP_BASE}/use-cases/iac-scan/"),
    ])
    doc.add_page_break()

    # ── 5.4 CWPP ──
    h2(doc, "5.4 CWPP — Cloud Workload Protection Platform")
    body(doc,
        "AccuKnox CWPP provides eBPF-based runtime security for containers, VMs, and bare-metal, "
        "using KubeArmor for inline Zero Trust policy enforcement."
    )
    h3(doc, "How It Works — Architecture")
    body(doc,
        "KubeArmor deploys as a DaemonSet on Kubernetes worker nodes (or systemd service on VMs). "
        "It uses eBPF probes and BPF LSM hooks to observe process, file, and network activity at the "
        "kernel level. Based on auto-discovered behavioral baselines, it generates and enforces least-privilege "
        "policies inline — blocking unauthorized operations before damage occurs. Enforcement modes include "
        "Observe (audit only), Audit (log violations), and Enforce (block)."
    )
    h3(doc, "KubeArmor Support Matrix")
    body(doc, "KubeArmor supports K8s orchestrated workloads (DaemonSet) and VM/Bare-Metal workloads (systemd mode):")
    for b in [
        "K8s: EKS, AKS, GKE, OpenShift, Rancher, k3s, MicroK8s, on-prem, bare-metal, edge",
        "VM/Bare-Metal: Ubuntu, RHEL, CentOS, Amazon Linux, SUSE, Debian, Oracle Linux",
        "Supported LSMs: BPF-LSM, AppArmor, SELinux",
        "Container runtimes: Docker, containerd, CRI-O",
    ]:
        bullet(doc, b)
    h3(doc, "Key Capabilities")
    for b in [
        "Inline attack prevention at kernel level (not detect-and-respond)",
        "Automated Zero Trust policy discovery and generation",
        "Application micro-segmentation at pod level",
        "Process, file, and network allowlisting",
        "Hardening policies based on MITRE, NIST, CIS, PCI-DSS, STIGs",
        "Container forensics and runtime telemetry",
        "Custom policy editor (observe/audit/enforce modes)",
    ]:
        bullet(doc, b)
    h3(doc, "Use Cases")
    bullet(doc, "Blocking cryptocurrency mining processes in containerized workloads using process allowlists", bold_prefix="Cryptojacking Prevention: ")
    bullet(doc, "Discovering application behavior and auto-generating Zero Trust policies that restrict file/process/network access to only observed patterns", bold_prefix="Zero Trust Policy Automation: ")
    resource_block(doc, [
        ("CWPP Overview", f"{HELP_BASE}/use-cases/cwpp/"),
        ("Runtime Security Architecture", f"{HELP_BASE}/getting-started/runtime-sec-arch/"),
        ("KubeArmor Support Matrix", f"{HELP_BASE}/support-matrix/kubearmor-support-matrix/"),
        ("Zero Trust Use Case", f"{HELP_BASE}/use-cases/zero-trust/"),
        ("Cryptojacking Prevention", f"{HELP_BASE}/use-cases/crypto-mining/"),
    ])

    # ── 5.5 CDR ── (NEW)
    h2(doc, "5.5 CDR — Cloud Detection & Response")
    body(doc,
        "AccuKnox CDR provides real-time cloud threat detection and automated remediation "
        "by ingesting cloud audit logs, applying detection rules, and triggering response workflows."
    )
    h3(doc, "How It Works — Architecture")
    body(doc,
        "CDR follows an event-driven pipeline: (1) Cloud logs (AWS CloudTrail, Azure Activity Logs, GCP Audit Logs) "
        "are ingested via Lambda/Cloud Functions into AccuKnox SIEM. (2) Detection rules engine evaluates events "
        "against security policies. (3) On violation, CDR triggers automated remediation via GitHub Actions/Jenkins "
        "webhooks using the accuknox/cdr-remediation action. (4) Notifications are sent via Slack, Jira, or Email. "
        "The remediation loop is fully automated and closed — detect, alert, fix, verify."
    )
    h3(doc, "Cloud Support")
    add_styled_table(doc,
        ["Cloud", "Log Source", "Remediation Examples"],
        [
            ["AWS", "CloudTrail → Lambda → SIEM", "Make S3 bucket private, shutdown public EC2 instance, start CloudTrail logging"],
            ["Azure", "Activity Logs → Event Hub → SIEM", "Block public storage access, terminate VM with public IP, enforce NSG rules"],
            ["GCP", "Audit Logs → Cloud Function → SIEM", "Revert firewall ports, delete instance with public IP, enforce IAM policies"],
        ],
        col_widths=[2.5, 5, 9.5]
    )
    h3(doc, "Key Capabilities")
    for b in [
        "Real-time cloud log ingestion and analysis via AccuKnox SIEM",
        "Policy-based detection engine with customizable rules",
        "Automated remediation via GitHub Actions / Jenkins / Webhook workflows",
        "Closed-loop response: Detect → Alert → Remediate → Verify",
        "Notification integration with Slack, Jira, Email",
        "Supports AWS CloudFormation–based deployment for CDR infrastructure",
    ]:
        bullet(doc, b)
    h3(doc, "Use Cases")
    bullet(doc, "Auto-revoking public access on S3 buckets within minutes of misconfiguration via CDR remediation pipeline", bold_prefix="Enforcing Private S3 Buckets: ")
    bullet(doc, "Alerting and auto-remediating when API calls originate from unapproved geographic regions", bold_prefix="Geo-Fencing Access Violations: ")
    resource_block(doc, [
        ("CDR Overview", f"{HELP_BASE}/use-cases/cdr/"),
        ("AWS CDR Onboarding", f"{HELP_BASE}/getting-started/aws-cdr/"),
        ("Azure CDR Onboarding", f"{HELP_BASE}/getting-started/azure-cdr/"),
        ("GCP CDR Onboarding", f"{HELP_BASE}/getting-started/gcp-cdr/"),
        ("CDR Remediation Setup", f"{HELP_BASE}/getting-started/cdr-setup/"),
    ])
    doc.add_page_break()

    # ── 5.6 API Security ──
    h2(doc, "5.6 API Security")
    body(doc,
        "AccuKnox API Security provides continuous API discovery, inventory, and protection "
        "across Kubernetes and microservices environments."
    )
    h3(doc, "How It Works")
    body(doc,
        "Runtime API Security uses service mesh sidecars or proxies (Istio, Nginx Ingress, Kong, F5) to inspect "
        "traffic and detect anomalies, exporting data in OpenTelemetry format. Static API Security scans "
        "code repos and API specs (OpenAPI, Swagger, WSDL). Combined, these provide full lifecycle API protection."
    )
    h3(doc, "Key Capabilities")
    for b in [
        "Real-time API endpoint inventory with method, path, sensitive data classification",
        "Shadow, zombie, and orphan API detection",
        "OWASP API Top 10 vulnerability coverage",
        "North-south and east-west traffic visibility and segmentation",
        "PII/PHI sensitive data detection in headers and responses",
        "Rate limiting: per-user, per-IP, per-endpoint, and global quotas",
        "Rate limit actions: Block (temporary/permanent) or Alert (log only)",
        "Access policy control and behavioral analytics",
        "Traffic connectors: AWS API Gateway, Kubernetes API Proxy, Istio, Nginx, Kong, F5",
    ]:
        bullet(doc, b)
    h3(doc, "Use Cases")
    bullet(doc, "Discovering undocumented shadow APIs across microservices and mapping them to OWASP API Top 10 risks", bold_prefix="Shadow API Discovery: ")
    bullet(doc, "Applying per-user rate limits on public-facing API endpoints to prevent abuse and DDoS", bold_prefix="API Rate Limiting: ")
    resource_block(doc, [
        ("API Security Use Case", f"{HELP_BASE}/use-cases/api-security/"),
        ("API Security Overview", f"{HELP_BASE}/integrations/api-overview/"),
        ("Kubernetes API Proxy", f"{HELP_BASE}/integrations/api-k8s/"),
    ])

    # ── 5.7 AI Security ──
    h2(doc, "5.7 AI Security")
    body(doc,
        "AccuKnox AI Security addresses AI-specific threats across models, agents, datasets, "
        "and pipelines with posture management, red teaming, and runtime protection."
    )
    h3(doc, "Architecture — Key Components")
    body(doc,
        "Prompt Firewall: Sits between the user/application and the LLM. Inspects prompts and responses "
        "in real-time using LLM-as-judge, regex-based PII masking, and policy checks. Supports SDK integration "
        "(Python/JS), OpenAI browser plugin, and AI gateway methods (Azure APIM, AWS API Gateway, LiteLLM, BiFrost)."
    )
    body(doc,
        "AI-DR (AI Detection & Response): Monitors AI/ML control-plane activity (CloudTrail, Azure Event Hub) "
        "out-of-band. Detects high-risk events like unauthorized model creation, insecure notebook configs, "
        "unapproved fine-tuning, and resource deletion. Triggers automated alerts and remediation."
    )
    body(doc,
        "Red Teaming: Automated adversarial testing for prompt injection, hallucinations, toxicity, bias, "
        "jailbreak resilience across 100+ supported models."
    )
    h3(doc, "Prompt Firewall Policy Types")
    add_styled_table(doc,
        ["Policy Type", "Purpose"],
        [
            ["Ban Code", "Prevent unauthorized code execution in prompts/responses"],
            ["Prompt Injection", "Guard against LLM manipulation and jailbreak attempts"],
            ["Toxicity", "Block harmful, hateful, or abusive language"],
            ["PII / Secrets", "Mask PII via regex patterns; prevent credential processing"],
            ["Relevance", "Keep input/output on-topic; filter off-topic content"],
            ["Sentiment", "Evaluate and flag user tone"],
            ["Token Limit", "Prevent DoS attacks and excessive cost via token caps"],
            ["Language", "Enforce approved languages for input/output"],
            ["Ban Topics / Competitors", "Restrict sensitive subjects and competitor mentions"],
        ],
        col_widths=[4, 13]
    )
    h3(doc, "AI/ML Support Matrix")
    add_styled_table(doc,
        ["Cloud Provider", "AI Services Supported"],
        [
            ["AWS", "SageMaker (Notebooks, Endpoints, Training), Bedrock (Model Customization, Invocation)"],
            ["Azure", "Azure ML Workspaces, Azure OpenAI Resources, Copilot Studio"],
            ["GCP", "Vertex AI, Cloud AI Platform"],
            ["On-Prem", "Hugging Face, Ollama, TensorFlow, PyTorch, private models"],
        ],
        col_widths=[4, 13]
    )
    h3(doc, "Use Cases")
    bullet(doc, "Blocking prompt injection and PII leakage in customer-facing chatbots using the Prompt Firewall SDK", bold_prefix="Prompt Firewall: ")
    bullet(doc, "Detecting unauthorized SageMaker notebook creation with public internet access and auto-alerting via AI-DR", bold_prefix="AI-DR: ")
    resource_block(doc, [
        ("AI/ML Overview", f"{HELP_BASE}/how-to/aiml-overview/"),
        ("Prompt Firewall Use Case", f"{HELP_BASE}/use-cases/prompt-firewall/"),
        ("AI-DR Use Case", f"{HELP_BASE}/use-cases/aidr/"),
        ("Azure AI-DR", f"{HELP_BASE}/use-cases/azure-aidr/"),
        ("LLM Defense SDK", f"{HELP_BASE}/use-cases/llm-defense-app-onboard/"),
        ("AI/ML Support Matrix", f"{HELP_BASE}/support-matrix/aiml-support-matrix/"),
        ("Copilot Studio Integration", f"{HELP_BASE}/integrations/copilot-studio/"),
    ])
    doc.add_page_break()

    # ── 5.8 Secrets Mgmt ──
    h2(doc, "5.8 Secrets Management")
    body(doc,
        "AccuKnox Secrets Management is a centralized, managed solution for credential security "
        "that serves as a drop-in replacement for HashiCorp Vault."
    )
    h3(doc, "Key Capabilities")
    for b in [
        "Encrypted secrets storage with versioning",
        "Dynamic secrets for temporary access to AWS, Kubernetes, databases",
        "Data encryption as a service (PKI, Transit encryption via API)",
        "Identity-based auth (LDAP, OIDC, OKTA) with granular permissions",
        "Full audit logging — direct integration with AccuKnox SIEM & CDR",
        "Multi-tenancy with per-tenant namespaces",
        "OS-level hardening powered by AccuKnox CWPP",
        "Deployable on-prem and in air-gapped environments",
    ]:
        bullet(doc, b)
    h3(doc, "Use Cases")
    bullet(doc, "Replacing HashiCorp Vault Enterprise with AccuKnox at lower cost while maintaining API compatibility", bold_prefix="Vault Migration: ")
    bullet(doc, "Generating short-lived AWS credentials for CI/CD pipelines via dynamic secrets", bold_prefix="Dynamic Secrets for CI/CD: ")
    resource_block(doc, [
        ("Secrets Management Guide", f"{HELP_BASE}/getting-started/secrets-management/"),
    ])

    # ── 5.9 SIEM ──
    h2(doc, "5.9 SIEM")
    body(doc,
        "AccuKnox SIEM is a cloud-native, AI-first security information and event management "
        "platform built for modern SOC teams."
    )
    h3(doc, "Key Capabilities")
    for b in [
        "AI-driven threat detection with ML-based correlation (up to 80% noise reduction)",
        "Centralized log aggregation from Kubernetes, containers, cloud services, endpoints",
        "Flexible ingestion: Syslog, KubeArmor, CloudTrail, Azure Logs, threat intel feeds",
        "Pre-built compliance reporting (SOC 2, PCI, HIPAA, GDPR)",
        "SOC-optimized dark-mode UI to reduce analyst fatigue",
        "SOAR integration for automated incident response playbooks",
        "10,000+ events/second processing, 100GB+ daily ingestion, sub-second search",
        "Hot/warm/cold data tiering for cost efficiency",
        "Built on OpenSearch for scalable analytics",
    ]:
        bullet(doc, b)
    h3(doc, "Use Cases")
    bullet(doc, "Correlating KubeArmor runtime alerts with CloudTrail events to surface multi-stage attack chains", bold_prefix="Cross-Source Threat Correlation: ")
    bullet(doc, "Replacing legacy Splunk/QRadar with AccuKnox SIEM for 40-60% MTTR reduction and predictable pricing", bold_prefix="Legacy SIEM Migration: ")

    doc.add_page_break()

    # ── 5.10 VM Security (expanded) ── (UPDATED)
    h2(doc, "5.10 VM Security — Vulnerability & Malware Scanning")
    body(doc,
        "AccuKnox VM Security provides comprehensive vulnerability scanning, malware detection, "
        "compliance benchmarking, and runtime hardening for virtual machines and bare-metal hosts."
    )
    h3(doc, "How It Works")
    body(doc,
        "Agentless scanning: Uses Terraform scripts to create point-in-time VM disk snapshots, "
        "which are analyzed for vulnerabilities and misconfigurations without installing any agent on the host. "
        "Supports AWS, Azure, GCP, and OpenShift Virtualization."
    )
    body(doc,
        "Agent-based scanning: The knoxctl CLI agent is deployed on Linux/Windows hosts to perform "
        "file-level vulnerability scanning with continuous monitoring. Supports air-gapped environments."
    )
    h3(doc, "VM Support Matrix")
    add_styled_table(doc,
        ["OS Family", "Distributions Supported", "Package Managers"],
        [
            ["RHEL-based", "RHEL 6-9, CentOS 6-8, AlmaLinux 8-10, Rocky 8-9, Oracle Linux 5-8, Amazon Linux 1/2/2023", "dnf / yum / rpm"],
            ["Debian-based", "Debian 7-12, Ubuntu (all Canonical-supported)", "apt / dpkg"],
            ["SUSE-based", "openSUSE Leap 42/15, SLES 11-15, SUSE Micro 5-6", "zypper / rpm"],
            ["Alpine", "Alpine Linux 2.2-3.22, edge", "apk"],
            ["Other Linux", "Photon OS 1.0-5.0, Azure Linux (CBL-Mariner) 1.0-3.0, Bottlerocket 1.7.0+", "tdnf / rpm / bottlerocket"],
            ["Windows", "Windows Server (agent-based via knoxctl)", "—"],
        ],
        col_widths=[3, 9, 5]
    )
    h3(doc, "Malware Scanning (ClamAV)")
    body(doc,
        "AccuKnox integrates ClamAV (Cisco's open-source antivirus engine) with 8M+ virus signatures "
        "maintained by Cisco Talos for comprehensive malware detection:"
    )
    for b in [
        "Quick scan for on-demand file/directory scanning with include/exclude and quarantine from SaaS UI",
        "Real-time protection on Linux via ClamOnAcc + ClamD on-access scanning daemon",
        "Archive scanning: RAR, 7Zip, Zip, Tar, XZ, Gzip, Bzip2, XAR, ARJ, IMG, ISO 9660, PKG",
        "Windows PE scanning: 32/64-bit executables with support for packed binaries (UPX, PeSpin, AsPack, etc.)",
        "Mail attachment scanning across all common mail formats",
        "PII detection: Credit cards (VISA, MasterCard, AMEX, Discover, JCB), U.S. SSNs",
    ]:
        bullet(doc, b)
    h3(doc, "Use Cases")
    bullet(doc, "Detecting CVE-2021-44228 (Log4j) across a fleet of cloud VMs using agentless scanning with auto-generated remediation guidance", bold_prefix="Log4Shell Detection: ")
    bullet(doc, "Continuous CIS/STIG compliance benchmarking for VM hosts with drift alerting", bold_prefix="Compliance Benchmarking: ")
    resource_block(doc, [
        ("VM Overview", f"{HELP_BASE}/use-cases/vm-overview/"),
        ("VM Host Scan", f"{HELP_BASE}/use-cases/vm-host-scan/"),
        ("VM Malware Scan", f"{HELP_BASE}/use-cases/vm-malware-scan/"),
        ("VM Compliance Benchmarking", f"{HELP_BASE}/use-cases/vm-compliance-benchmarking/"),
        ("VM Support Matrix", f"{HELP_BASE}/support-matrix/vms/"),
        ("Agent-Based Linux", f"{HELP_BASE}/how-to/vm-security/agent-based/linux/"),
        ("Agent-Based Windows", f"{HELP_BASE}/how-to/vm-security/agent-based/windows/"),
        ("Agentless Cloud VM Scan", f"{HELP_BASE}/how-to/vm-security/agentless/cloud-vm-scanning/"),
    ])

    # ── 5.11 Container Registry Scanning ── (NEW)
    h2(doc, "5.11 Container Registry Scanning")
    body(doc,
        "AccuKnox provides vulnerability, sensitive data, and layer analysis for container images "
        "across 10 supported registries, available via both SaaS and on-prem/air-gapped deployments."
    )
    h3(doc, "Registry Support Matrix")
    add_styled_table(doc,
        ["Registry", "Auth Method", "Deployment Types"],
        [
            ["AWS ECR", "IAM-based", "Cloud, Hybrid"],
            ["Azure ACR", "Basic Auth", "Cloud, Hybrid"],
            ["Google Artifact Registry", "Service Account", "Cloud, Hybrid"],
            ["Docker Hub", "Basic Auth", "Cloud, On-Prem"],
            ["Harbor", "Basic Auth", "On-Prem, Hybrid"],
            ["JFrog", "Basic Auth", "Cloud, On-Prem, Hybrid"],
            ["Quay", "Basic Auth", "Cloud, On-Prem"],
            ["Sonatype Nexus", "Basic Auth", "On-Prem, Hybrid"],
            ["Docker Trusted Registry", "Basic Auth", "On-Prem"],
            ["In-Cluster Scanner", "K8s RBAC", "On-Prem, Air-Gapped"],
        ],
        col_widths=[4.5, 4, 8.5]
    )
    h3(doc, "Key Capabilities")
    for b in [
        "Automated and on-demand container image scanning",
        "Dashboard views: Vulnerabilities (CVEs + severity), Sensitive Data (secrets/credentials), Resources (libraries), Layers",
        "ECR automated scan support with event-driven triggers",
        "In-cluster scanner for air-gapped / self-hosted environments",
        "CI/CD integration for container scanning in build pipelines",
        "Ticket creation integration (Jira, ServiceNow) for vulnerability findings",
    ]:
        bullet(doc, b)
    h3(doc, "Use Cases")
    bullet(doc, "Scanning all images in an ECR registry on every push and blocking deployments with critical CVEs", bold_prefix="Automated Registry Scanning: ")
    bullet(doc, "Running in-cluster scans in air-gapped environments where external registry access is restricted", bold_prefix="Air-Gapped Scanning: ")
    resource_block(doc, [
        ("Registry Overview", f"{HELP_BASE}/how-to/registry-overview/"),
        ("Registry Support Matrix", f"{HELP_BASE}/support-matrix/registry/"),
        ("ECR Automated Scan", f"{HELP_BASE}/how-to/ecr-automated-scan/"),
        ("In-Cluster Scanner", f"{HELP_BASE}/how-to/in-cluster-image-scan-helm/"),
        ("Container Scan Use Case", f"{HELP_BASE}/use-cases/container-scan/"),
        ("Container Image Scanning (Agent)", f"{HELP_BASE}/how-to/vm-security/agent-based/container-image-scanning/"),
    ])
    doc.add_page_break()


# ── Section 6: Integrations ──────────────────────────────────────────────────

def add_section_6(doc):
    h1(doc, "6. Integrations")
    body(doc,
        "AccuKnox supports 30+ integrations enabling bi-directional sharing of security findings "
        "across the cloud security ecosystem."
    )

    h2(doc, "6.1 Notification Integrations")
    add_styled_table(doc,
        ["Platform", "Description", "Help Docs"],
        [
            ["Slack", "Real-time alerts to Slack channels", f"{HELP_BASE}/integrations/slack/"],
            ["Microsoft Teams", "Alert delivery for SOC and DevOps collaboration", "—"],
            ["PagerDuty", "Incident escalation and on-call routing", "—"],
            ["Email (SMTP)", "Configurable email notifications for findings", f"{HELP_BASE}/integrations/email/"],
            ["Webhook", "Generic webhook for custom notification workflows", f"{HELP_BASE}/integrations/webhook-integration/"],
        ],
        col_widths=[3.5, 6, 7.5]
    )

    h2(doc, "6.2 Ticketing Integrations")
    add_styled_table(doc,
        ["Platform", "Description", "Help Docs"],
        [
            ["Jira Cloud", "Auto ticket creation with context, severity, remediation", f"{HELP_BASE}/integrations/jira-cloud/"],
            ["Jira Server", "On-prem Jira integration for CSPM/CWPP findings", f"{HELP_BASE}/integrations/jira-server-cspm/"],
            ["ServiceNow", "ITSM integration for incident management", f"{HELP_BASE}/integrations/servicenow/"],
            ["Freshservice", "CSPM findings integration with Freshservice", f"{HELP_BASE}/integrations/freshservice-cspm/"],
            ["Connectwise", "CSPM findings integration with Connectwise", f"{HELP_BASE}/integrations/connectwise-cspm/"],
            ["ServiceDesk Plus", "Ticket creation and management", f"{HELP_BASE}/integrations/servicedesk-plus/"],
        ],
        col_widths=[3.5, 6, 7.5]
    )

    h2(doc, "6.3 SIEM & Log Management")
    add_styled_table(doc,
        ["Platform", "Description", "Help Docs"],
        [
            ["Splunk", "Log forwarding and security events via HEC", f"{HELP_BASE}/integrations/splunk/"],
            ["IBM QRadar", "SIEM event integration", f"{HELP_BASE}/integrations/ibm-qradar/"],
            ["Azure Sentinel", "Native Azure Sentinel integration", f"{HELP_BASE}/integrations/azure-sentinel/"],
            ["Rsyslog", "Standard syslog forwarding", f"{HELP_BASE}/integrations/rsyslog/"],
            ["AWS CloudWatch", "Cloud log integration", f"{HELP_BASE}/integrations/aws-cloudwatch/"],
            ["AccuKnox SIEM", "Built-in SIEM with native ingestion from all modules", f"{WEB_BASE}/platform/siem"],
        ],
        col_widths=[3.5, 6, 7.5]
    )

    h2(doc, "6.4 SSO Integrations")
    add_styled_table(doc,
        ["Platform", "Help Docs"],
        [
            ["Azure Entra ID", f"{HELP_BASE}/integrations/azure-entra-sso/"],
            ["Okta", f"{HELP_BASE}/integrations/okta-sso/"],
            ["Auth0", f"{HELP_BASE}/integrations/auth0-sso/"],
        ],
        col_widths=[5, 12]
    )
    doc.add_page_break()


# ── Section 7: Workflow Automation ───────────────────────────────────────────

def add_section_7(doc):
    h1(doc, "7. Workflow Automation")

    h2(doc, "7.1 Rules Engine")
    body(doc,
        "The AccuKnox Rules Engine enables automated policy enforcement and response workflows "
        "across all security modules."
    )
    for b in [
        "Configurable rules based on severity, resource type, compliance framework, or custom attributes",
        "Automated actions: create tickets, send notifications, trigger remediation playbooks, suppress findings",
        "Multi-condition rules with AND/OR logic for precise targeting",
        "Rule templates for common automation scenarios",
        "Audit trail of all rule executions and outcomes",
    ]:
        bullet(doc, b)
    resource_block(doc, [
        ("Rules Engine Use Case", f"{HELP_BASE}/use-cases/rules-engine-ticket-creation/"),
        ("Findings Lifecycle", f"{HELP_BASE}/how-to/findings-lifecycle/"),
    ])

    h2(doc, "7.2 Ask AI Assistance (Ada)")
    body(doc,
        "AccuKnox provides a Gen-AI powered security assistant (Ada) that helps users investigate "
        "findings, understand risks, and accelerate remediation."
    )
    for b in [
        "Natural language queries about security findings, policies, and compliance status",
        "AI-generated remediation guidance with step-by-step instructions",
        "Context-aware analysis that correlates findings across modules",
        "Interactive exploration of cloud attack paths and risk relationships",
    ]:
        bullet(doc, b)
    doc.add_page_break()


# ── Section 8: SLA, Support, Resources ───────────────────────────────────────

def add_section_8(doc):
    h1(doc, "8. SLA, Support Coverage, and Resources")

    h2(doc, "8.1 Support Tiers")
    add_styled_table(doc,
        ["Support Level", "Support Hours"],
        [
            ["Silver (Standard)", "12x5 (6:00 AM – 6:00 PM PT, Mon-Fri, excluding US holidays)"],
            ["Gold (Premium)", "12x7 (Extended Phone Support)"],
            ["Platinum (PremiumPlus)", "24x7 (Full Round-the-Clock Support)"],
        ],
        col_widths=[5, 12]
    )

    h2(doc, "8.2 Additional Support Features")
    add_styled_table(doc,
        ["Feature", "Silver (Standard)", "Gold (Premium)", "Platinum (PremiumPlus)"],
        [
            ["Training", "Standard Training", "Standard Training + 2 Student Admin Webinar Series",
             "Live Training + 5 Student Premium Training Enrollments"],
            ["Phone Support", "12x5", "12x7", "24x7"],
            ["Online Support", "✓", "✓", "✓"],
            ["Adoption Services", "None", "None", "Named Customer Success Manager"],
        ],
        col_widths=[4, 4, 4.5, 4.5]
    )

    h2(doc, "8.3 SLA Response Times")
    add_styled_table(doc,
        ["Severity Level", "Silver (Standard)", "Gold (Premium)", "Platinum (PremiumPlus)"],
        [
            ["P1 (Critical)", "First Response: 4 Hours\nUpdates: 1 Business Day",
             "First Response: 1 Hour\nUpdates: 2 Hours", "First Response: 1 Hour\nUpdates: 2 Hours"],
            ["P2 (High)", "First Response: 1 Business Day\nUpdates: 2 Business Days",
             "First Response: 2 Hours\nUpdates: 8 Hours", "First Response: 2 Hours\nUpdates: 8 Hours"],
            ["P3 (Medium)", "First Response: 2 Business Days\nUpdates: 3 Business Days",
             "First Response: 2 Hours\nUpdates: 48 Hours", "First Response: 2 Hours\nUpdates: 48 Hours"],
            ["P4 (Low)", "First Response: 2 Business Days\nUpdates: 3 Business Days",
             "First Response: 8 Hours\nUpdates: 48 Hours", "First Response: 8 Hours\nUpdates: 48 Hours"],
        ],
        col_widths=[3.5, 4.5, 4.5, 4.5]
    )
    doc.add_page_break()

    h2(doc, "8.4 Severity Levels & Definitions")
    add_styled_table(doc,
        ["Severity Level", "Definition", "Examples"],
        [
            ["P1 (Critical)", "AccuKnox is down and inaccessible. Severe service failure affecting multiple users.",
             "Users cannot access the platform. Consistent errors prevent login."],
            ["P2 (High)", "Partial service failure or mild degradation. Some resources accessible.",
             "Admin console write-access issue. Slow access, occasional errors. Significant impact to integration."],
            ["P3 (Medium)", "Minor service impact, affecting individual users or non-critical apps.",
             "One user unable to access an application. Difficulty integrating new applications."],
            ["P4 (Low)", "Minor impact or feature enhancement request.",
             "How-to inquiries. Feature enhancement requests."],
        ],
        col_widths=[3, 6.5, 7.5]
    )

    h2(doc, "8.5 Escalation Matrix")
    add_styled_table(doc,
        ["Severity Level", "Personnel To Contact", "Email"],
        [
            ["L1 (Low)", "Udit, Aditya — Sr Solution Engineers @AccuKnox", "udit@accuknox.com, adityaraj@accuknox.com"],
            ["L2 (Medium)", "Gaurav — Product Manager @AccuKnox", "gaurav.mishra@accuknox.com"],
            ["L3 (Critical)", "Rahul Jadhav — CTO @AccuKnox", "r@accuknox.com"],
        ],
        col_widths=[3, 6.5, 7.5]
    )

    h2(doc, "8.6 Contact Methods")
    body(doc, "You can reach our support team through the following methods, listed in order of preference:")
    bullet(doc, "Send an email directly to support@accuknox.com", bold_prefix="Email: ")
    bullet(doc, "Raise a support ticket via our customer portal: https://accu-knox.atlassian.net/servicedesk/customer/portal/1", bold_prefix="Jira Service Desk: ")
    body(doc, "Our support team will respond to your ticket within 24 business hours.")

    h2(doc, "8.7 Support Resources")
    add_styled_table(doc,
        ["Resource", "URL"],
        [
            ["Support Email", "support@accuknox.com"],
            ["Help Portal", "https://help.accuknox.com"],
            ["Status Page", "https://status.accuknox.com"],
            ["Jira Service Desk", "https://accu-knox.atlassian.net/servicedesk/customer/portal/1"],
            ["SLA & Escalation Matrix", f"{HELP_BASE}/resources/sla-escalation-matrix/"],
            ["Technical Support Guide", f"{HELP_BASE}/resources/technical-support-guide/"],
            ["Ticketing Procedures", f"{HELP_BASE}/resources/ticket-procedure/"],
        ],
        col_widths=[5, 12]
    )

    h2(doc, "8.8 Recovery Objectives")
    add_styled_table(doc,
        ["Metric", "Commitment", "Scope"],
        [
            ["RTO", "6 hours", "Core control plane services (policy management, telemetry, dashboards, enforcement)."],
            ["RPO", "24 hours", "Control plane config, security policies, tenant config, audit logs, platform state."],
        ],
        col_widths=[3, 3, 11]
    )
    body(doc,
        "Note: Runtime enforcement on customer workloads continues independently via deployed agents, "
        "ensuring fail-secure behavior even during control plane recovery."
    )


# ── Footer ────────────────────────────────────────────────────────────────────

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("AccuKnox Inc. | CONFIDENTIAL | support@accuknox.com | accuknox.com")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_cover_page(doc)
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)
    add_section_4(doc)
    add_section_5(doc)
    add_section_6(doc)
    add_section_7(doc)
    add_section_8(doc)
    add_footer(doc)

    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
